"""Build the concise classroom handout for the portfolio project."""

from __future__ import annotations

import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "courseware" / "AI图像识别项目课堂讲义.docx"
FONT_REGULAR = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD = Path("C:/Windows/Fonts/msyhbd.ttc")
BLUE = "2457A6"
LIGHT_BLUE = "EAF1FB"
PALE_BLUE = "F5F8FC"
GRAY = "666666"
LIGHT_GRAY = "D9E0E8"
BLACK = "000000"


def image_font(size: int, bold: bool = False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LIGHT_GRAY, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(BLACK)
    title.paragraph_format.space_after = Pt(14)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    for name, size, before, after in (
        ("Heading 1", 17, 18, 8),
        ("Heading 2", 13, 14, 6),
        ("Heading 3", 11, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("AI Image Recognition 课堂讲义")
        set_run_font(run, 8.5, color=GRAY)
        add_page_number(sec.footer.paragraphs[0])
    return doc


def add_paragraph(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(item))


def add_numbered(doc: Document, items: list[str]) -> None:
    for number, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(f"{number}.  {item}"))


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F7")
    p_pr.append(shd)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = Cm(width)
        set_cell_shading(cell, BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        set_run_font(paragraph.add_run(header), 9.5, True, "FFFFFF")

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(row_values, widths)):
            cell = cells[index]
            cell.width = Cm(width)
            set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else PALE_BLUE)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.keep_with_next = row_index < len(rows) - 1
            if index == 0 and len(value) < 20:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(value), 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def draw_flow(path: Path, title: str, nodes: list[tuple[float, float, float, float, str, str]], arrows: list[tuple[int, int]], subtitle: str) -> None:
    width, height = 1920, 820
    canvas = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = image_font(42, bold=True)
    node_font = image_font(29, bold=True)
    subtitle_font = image_font(23)

    def px(node):
        x, y, w, h, label, color = node
        # Diagram coordinates use a bottom-left origin; Pillow uses top-left.
        left = int(x * width)
        top = int((1 - y - h) * height)
        right = int((x + w) * width)
        bottom = int((1 - y) * height)
        return left, top, right, bottom, label, color

    boxes = [px(node) for node in nodes]
    for left, top, right, bottom, label, color in boxes:
        draw.rounded_rectangle(
            (left, top, right, bottom),
            radius=14,
            fill="#FFFFFF",
            outline=color,
            width=4,
        )
        box = draw.multiline_textbbox((0, 0), label, font=node_font, spacing=7, align="center")
        text_width = box[2] - box[0]
        text_height = box[3] - box[1]
        draw.multiline_text(
            ((left + right - text_width) / 2, (top + bottom - text_height) / 2 - 3),
            label,
            fill="#152238",
            font=node_font,
            spacing=7,
            align="center",
        )

    for start, end in arrows:
        sl, st, sr, sb, _, _ = boxes[start]
        el, et, er, eb, _, _ = boxes[end]
        if el > sr:
            p1 = (sr + 8, (st + sb) // 2)
            p2 = (el - 14, (et + eb) // 2)
        elif er < sl:
            p1 = (sl - 8, (st + sb) // 2)
            p2 = (er + 14, (et + eb) // 2)
        elif et > sb:
            p1 = ((sl + sr) // 2, sb + 8)
            p2 = ((el + er) // 2, et - 14)
        else:
            p1 = ((sl + sr) // 2, st - 8)
            p2 = ((el + er) // 2, eb + 14)
        draw.line((p1, p2), fill="#6B778C", width=4)
        x2, y2 = p2
        if abs(p2[0] - p1[0]) >= abs(p2[1] - p1[1]):
            direction = 1 if p2[0] > p1[0] else -1
            head = [(x2, y2), (x2 - 18 * direction, y2 - 10), (x2 - 18 * direction, y2 + 10)]
        else:
            direction = 1 if p2[1] > p1[1] else -1
            head = [(x2, y2), (x2 - 10, y2 - 18 * direction), (x2 + 10, y2 - 18 * direction)]
        draw.polygon(head, fill="#6B778C")

    draw.text((38, 24), title, fill="#000000", font=title_font)
    draw.text((38, height - 54), subtitle, fill="#5D6675", font=subtitle_font)
    canvas.save(path, format="PNG", optimize=True)


def make_figures(folder: Path) -> dict[str, Path]:
    figures: dict[str, Path] = {}

    architecture = folder / "architecture.png"
    draw_flow(
        architecture,
        "系统整体架构",
        [
            (0.03, 0.40, 0.15, 0.20, "浏览器前端\n上传与展示", "#2457A6"),
            (0.25, 0.40, 0.15, 0.20, "Flask 后端\n校验与编排", "#2457A6"),
            (0.48, 0.63, 0.18, 0.17, "MobileNetV3\n37 类分类", "#2A7F62"),
            (0.48, 0.20, 0.18, 0.17, "YOLOv8n\n通用目标检测", "#A25D21"),
            (0.74, 0.20, 0.18, 0.17, "OpenCV\n绘框与保存", "#A25D21"),
            (0.75, 0.63, 0.17, 0.17, "JSON 结果\nTop 3", "#2A7F62"),
        ],
        [(0, 1), (1, 2), (1, 3), (2, 5), (3, 4)],
        "前端负责交互，Flask 负责协调，模型负责推理，OpenCV 负责检测结果可视化。",
    )
    figures["architecture"] = architecture

    inference = folder / "inference.png"
    draw_flow(
        inference,
        "分类请求的数据流",
        [
            (0.02, 0.40, 0.13, 0.19, "上传图片", "#2457A6"),
            (0.19, 0.40, 0.15, 0.19, "格式与内容\n校验", "#2457A6"),
            (0.38, 0.40, 0.15, 0.19, "RGB 缩放\n裁剪归一化", "#2A7F62"),
            (0.57, 0.40, 0.15, 0.19, "MobileNetV3\n前向推理", "#2A7F62"),
            (0.76, 0.40, 0.18, 0.19, "Softmax 与 Top 3\n返回 JSON", "#2A7F62"),
        ],
        [(0, 1), (1, 2), (2, 3), (3, 4)],
        "训练和推理必须使用匹配的预处理；Top 3 是三个候选类别，不是三个目标。",
    )
    figures["inference"] = inference

    cpsc = folder / "cpsc210.png"
    draw_flow(
        cpsc,
        "CPSC 210 设计概念在项目中的落地",
        [
            (0.03, 0.63, 0.18, 0.17, "ModelLoader\n抽象接口", "#2457A6"),
            (0.29, 0.63, 0.21, 0.17, "FineTunedLoader\n具体实现", "#2A7F62"),
            (0.03, 0.20, 0.18, 0.17, "ObjectDetector\n抽象接口", "#2457A6"),
            (0.29, 0.20, 0.21, 0.17, "YOLOv8Detector\n具体实现", "#A25D21"),
            (0.59, 0.20, 0.18, 0.17, "DetectionResult\n不可变值对象", "#7A4EAB"),
            (0.78, 0.60, 0.17, 0.18, "web_app\n依赖抽象能力", "#2457A6"),
        ],
        [(0, 1), (2, 3), (3, 4), (1, 5), (4, 5)],
        "抽象规定能力，具体类完成实现，值对象稳定传递数据，Web 层只负责流程协调。",
    )
    figures["cpsc"] = cpsc
    return figures


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_run_font(cap.add_run(caption), 9, color=GRAY)


def build() -> None:
    with tempfile.TemporaryDirectory(prefix="ai_handout_") as temp_dir:
        figures = make_figures(Path(temp_dir))
        doc = setup_document()

        title = doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(72)
        set_run_font(title.add_run("AI 图像识别项目课堂讲义"), 26, True, BLACK)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(24)
        set_run_font(subtitle.add_run("从模型推理到 Web 作品的完整闭环"), 14, color=GRAY)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(26)
        set_run_font(meta.add_run("适用内容  第六课与第七课综合复习\n建议课时  60 至 90 分钟\n学习对象  具备基础 Python 知识的初学者"), 10.5, color=GRAY)

        intro = doc.add_paragraph()
        intro.paragraph_format.left_indent = Cm(1.0)
        intro.paragraph_format.right_indent = Cm(1.0)
        intro.paragraph_format.line_spacing = 1.45
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_run_font(intro.add_run("本讲义帮助同学从整体上理解项目，而不是逐行背代码。学习重点是看清数据、模型、后端、前端和测试怎样连接起来，并能够准确区分图像分类、目标检测、预训练、微调和图像处理。学完后，同学应当能够独立演示系统，并用 CPSC 210 的面向对象概念解释主要设计。"), 11)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        doc.add_heading("一  为什么要做这堂综合课", level=1)
        add_paragraph(doc, "前面的课程已经分别接触了数据集、迁移学习、模型微调、Flask 页面和 YOLO 检测。如果只把这些知识点分开记忆，同学容易知道每个名词，却不知道它们为什么要放在同一个项目中。这堂课的目标是把零散知识串成一条完整主线：用户上传图片，后端验证输入，根据任务调用合适模型，将模型输出转换成稳定数据，再由页面向用户展示。")
        add_paragraph(doc, "一个作品级项目不仅要让模型跑出结果，还要让别人能够安装、启动、操作、理解和排查。模型准确率只是质量的一部分；路径管理、依赖环境、错误提示、接口设计、临时文件清理和自动化测试同样决定作品是否可靠。")

        doc.add_heading("学习目标", level=2)
        add_bullets(doc, [
            "能用一句话说明项目解决的问题和面向的用户。",
            "能区分整图分类和目标检测，并说明为什么使用两个模型。",
            "能解释 Flask、MobileNetV3、YOLOv8 和 OpenCV 各自的职责。",
            "能说清 CPSC 210 中抽象、继承、多态、值对象和单一职责在代码中的位置。",
            "能按照规范步骤启动、演示和测试项目。",
            "能主动说明当前数据集与模型的能力边界。",
        ])

        doc.add_heading("二  项目概览与能力边界", level=1)
        add_paragraph(doc, "本项目是一个基于 PyTorch、TorchVision、Flask、YOLOv8 和 OpenCV 的图像识别 Web 系统。页面提供图像分类和目标检测两种模式。分类使用微调后的 MobileNetV3 Large，检测使用 YOLOv8n 通用预训练权重，Flask 负责接收图片和返回 JSON，OpenCV 负责生成带框结果图。")
        add_table(doc,
            ["项目部分", "当前实现", "必须说明的边界"],
            [
                ["图像分类", "Oxford IIIT Pet 37 类宠物品种，返回 Top 3", "不包含金渐层、银渐层、蓝猫等独立毛色标签"],
                ["目标检测", "YOLOv8n 通用物体检测，返回类别、坐标和置信度", "第七课没有重新训练 YOLO，也不是宠物品种检测器"],
                ["模型评估", "1,478 张验证图片，Top 1 为 90.19%，正确 1,333 张", "整体准确率不代表每个类别都达到相同水平"],
                ["Web 系统", "本地 Flask 页面、上传校验、结果展示和错误处理", "开发服务器适合课堂演示，不应直接用于公网生产"],
            ],
            [2.6, 7.1, 6.4],
        )
        add_paragraph(doc, "分类器是封闭集模型。即使输入一张猪或汽车图片，它仍会在已知 37 类中分配分数，因此高分不等于答案一定正确。若要识别新的毛色类别，必须先建立明确标签的数据集，再重新训练和评估，而不是只改网页上的类别文字。", "分类器是封闭集模型。")

        doc.add_heading("两个典型使用场景", level=2)
        add_paragraph(doc, "场景一是宠物图片品种判断。用户上传一张主体较清晰的猫狗照片，系统返回最相近的三个 Oxford IIIT Pet 品种。这个场景适合解释迁移学习、分类头、Softmax 和 Top 3，也适合展示分类模型在相似品种之间仍可能混淆。页面结果应被理解为模型建议，而不是宠物血统证明。")
        add_paragraph(doc, "场景二是复杂图片中的目标定位。用户上传一张同时包含人、狗、椅子或汽车的图片，YOLO 可能返回多个目标以及各自位置。这个场景适合解释检测框、置信度阈值和 OpenCV 绘图。检测器只知道预训练类别，能检测到 cat 或 dog 不代表它可以直接判断 Persian 或 Russian Blue。")
        add_paragraph(doc, "两个场景放在同一页面的教学价值，是让同学看到任务定义会决定数据、模型、输出和评估方式。分类数据只需要每张图的类别标签，检测训练还需要标出目标位置；分类常看 Accuracy 和 F1，检测还要考虑边界框重叠和 mAP。不能因为输入都是图片，就认为两种任务完全相同。")
        add_paragraph(doc, "如果未来要实现多只猫的毛色识别，可以先用检测器找到每只猫，再裁剪每个框并交给专门的毛色分类器。这个两阶段方案需要新的毛色标签、真实场景验证和未知类别处理。它展示了系统扩展应从任务和数据出发，而不是只把 YOLO、MobileNet 等模型名称堆在一起。")

        doc.add_heading("三  系统整体架构", level=1)
        add_figure(doc, figures["architecture"], "图 1  从浏览器到模型与结果展示的职责关系")
        add_paragraph(doc, "架构中最重要的是职责分离。前端只处理交互和展示，不读取模型权重；Flask 负责路由、输入校验和流程协调，不负责学习视觉特征；MobileNetV3 和 YOLO 负责推理；OpenCV 不判断物体类别，只根据检测结果绘图。这样的边界让每个模块更容易理解、测试和替换。")
        add_table(doc,
            ["文件", "主要职责", "与其他模块的联系"],
            [
                ["templates/index.html", "模式切换、拖拽上传、预览、状态和结果渲染", "通过 fetch 调用 predict 或 detect 接口"],
                ["app/web_app.py", "Flask 入口、路由、校验、模型缓存和 JSON 响应", "协调加载器、检测器、Pillow 与 OpenCV"],
                ["app/model_loader.py", "构建 MobileNetV3 并加载本地分类权重", "为 web_app 和评估脚本提供模型"],
                ["app/object_detector.py", "定义检测接口、YOLO 实现和 DetectionResult", "把 Ultralytics 输出转换成项目数据结构"],
                ["app/preprocess.py", "读取原图、绘制检测框和保存结果图", "消费 DetectionResult，不直接调用模型"],
            ],
            [4.2, 6.2, 5.7],
        )

        doc.add_heading("四  图像分类流程", level=1)
        add_figure(doc, figures["inference"], "图 2  图像分类从上传到 Top 3 的数据流")
        add_paragraph(doc, "分类任务回答的是整张图片最像哪个类别。浏览器把图片作为 multipart form data 发送给 `/predict`。后端先检查文件字段、扩展名和大小，再用 Pillow 验证它确实可以被解析为图片。文件保存时加入 UUID，避免不同用户上传同名文件造成覆盖。")
        add_paragraph(doc, "预处理阶段先把图片统一为 RGB，再缩放、中心裁剪到 224 乘 224、转换成 Tensor，并按照 ImageNet 的均值和标准差归一化。训练和推理预处理必须匹配，否则模型看到的数值分布发生变化，即使权重正确也可能明显掉点。")
        add_paragraph(doc, "MobileNetV3 输出 37 个 logits。logits 是未归一化数值，程序使用 Softmax 转成相对分数，再通过 Top K 选择最高的三个类别。Top 3 表示三个候选品种，不表示图中存在三个物体。模型实际类别只有两个时，Top 3 最多也只能返回两个结果。")

        doc.add_heading("迁移学习与微调", level=2)
        add_paragraph(doc, "ImageNet 预训练模型已经学习边缘、纹理和形状等通用视觉特征。项目将原来的分类头替换为 37 类输出，并使用 Oxford IIIT Pet 数据继续训练。数据较少时可以先冻结特征层，只更新分类头；需要提高任务适应性时，再使用较小学习率解冻更多层。当前正式 checkpoint 是从旧权重继续全模型微调得到的，使用 AdamW、权重衰减、label smoothing 和余弦学习率调度。这个过程不是 LoRA。")
        add_paragraph(doc, "推理时模型使用 `weights=None` 创建网络结构，然后加载项目中的完整 checkpoint，因此依赖安装完成后不需要再次下载 ImageNet 权重。模型调用 `eval()` 切换到推理行为，并在无梯度环境中前向计算，减少不必要的内存开销。")

        doc.add_heading("五  目标检测与 OpenCV", level=1)
        add_paragraph(doc, "检测任务不仅回答图里有什么，还要回答目标在哪里。YOLOv8Detector 从每个检测框中读取类别索引、类别名称、置信度和左上角及右下角坐标，并转换为项目统一的 DetectionResult。一次请求可以返回多个目标，也可以返回空列表。")
        add_table(doc,
            ["比较项", "图像分类", "目标检测"],
            [
                ["核心问题", "整张图片是什么", "图片中有什么并且在哪里"],
                ["当前模型", "微调 MobileNetV3 Large", "预训练 YOLOv8n"],
                ["主要输出", "37 类分数与 Top 3", "多个类别、置信度和边界框"],
                ["常见指标", "Accuracy、Precision、Recall、F1", "Precision、Recall、mAP"],
                ["本项目是否训练", "完成了宠物分类微调", "没有进行自定义 YOLO 微调"],
            ],
            [3.1, 6.5, 6.5],
        )
        add_paragraph(doc, "置信度阈值是推理阶段的过滤条件。默认值 0.5 在调用方没有提供值时生效；网页滑块可以把 0.3 或 0.7 发送给后端，并只覆盖本次请求。降低阈值通常保留更多框，同时可能增加误检；提高阈值会让结果更严格，同时可能造成漏检。调整阈值不会重新训练模型，也不会修改权重。")
        add_paragraph(doc, "OpenCV 在项目中负责读图、限制坐标范围、绘制矩形和文字、保存结果。YOLO 先提供哪里有什么，OpenCV 再把结构化结果画给用户看。没有 OpenCV 时仍可以返回坐标 JSON，但页面缺少直观的带框图片；只有 OpenCV 而没有检测模型，它并不知道哪里存在猫或狗。")

        doc.add_heading("六  Flask 与前端如何协作", level=1)
        add_paragraph(doc, "Flask 提供三个主要路由。`GET /` 返回页面并检查资源状态；`POST /predict` 接收分类图片；`POST /detect` 接收检测图片和阈值。模型采用懒加载：打开首页时不强制读取大权重，第一次真正推理时才加载，此后复用同一对象。这样既避免每次请求重新加载，也让模型缺失时首页仍能显示诊断信息。")
        add_table(doc,
            ["接口", "输入", "成功输出", "典型错误"],
            [
                ["GET /", "无", "HTML 与资源状态", "模板问题应通过绝对项目路径避免"],
                ["POST /predict", "file", "results、image_url、模型与数据路径", "400 无效图片，503 模型或数据缺失"],
                ["POST /detect", "file、conf_threshold", "detections、count、结果图 URL", "400 阈值错误，503 检测依赖不可用"],
            ],
            [2.5, 3.8, 6.2, 5.7],
        )
        add_paragraph(doc, "JSON 是前后端之间的稳定契约。分类结果包含 label 和 score，检测结果还包含坐标、count 和结果图片地址。前端根据这些字段生成列表、百分比条和图片。如果后端改变字段名称或分数单位，前端也必须同步修改，并通过接口测试防止契约回归。")
        add_paragraph(doc, "后端校验不能被前端校验替代。用户可以绕过页面直接访问接口，所以服务器仍要限制 16 MB、检查扩展名并验证真实图片内容。输入问题返回 400，文件过大返回 413，模型或依赖不可用返回 503，未预期异常记录日志后返回通用 500，避免把本机路径和堆栈直接暴露给用户。")

        doc.add_heading("七  与 CPSC 210 的结合", level=1)
        add_figure(doc, figures["cpsc"], "图 3  面向对象设计概念与项目代码的对应关系")
        add_paragraph(doc, "CPSC 210 关注的不只是语法，而是如何用抽象和对象组织可维护的软件。本项目中的深度学习模型可以被看成外部能力，面向对象设计负责规定这些能力如何被系统稳定使用。")

        doc.add_heading("抽象和面向接口编程", level=2)
        add_paragraph(doc, "`ModelLoader` 抽象类规定加载器必须提供 `load_model()`；`ObjectDetector` 规定检测器必须提供 `detect()`。抽象描述调用方需要的能力，不暴露某个库的全部内部细节。Web 层针对这些稳定能力组织流程，而不是把 Ultralytics 结果对象散布到每个页面函数中。")

        doc.add_heading("继承和多态", level=2)
        add_paragraph(doc, "`FineTunedLoader` 是 `ModelLoader` 的具体实现，`YOLOv8Detector` 是 `ObjectDetector` 的具体实现。不同实现可以共享同一个方法名和调用方式，这就是多态带来的替换能力。测试时可以放入 Fake 或 Mock 实现，让接口测试不必每次加载真实大模型。")

        doc.add_heading("值对象", level=2)
        add_paragraph(doc, "`DetectionResult` 是冻结的 dataclass，保存 x1、y1、x2、y2、label 和 confidence。它不负责推理，只负责稳定传递一次检测的结果。`frozen=True` 减少数据在模块之间被意外修改，`to_dict()` 则把同一份值转换成 JSON 所需格式。定义放在文件前部，是因为后面的检测器、绘图函数和 Web 层都要依赖这个类型；类必须先定义，后续代码才能创建它。")

        doc.add_heading("单一职责与低耦合", level=2)
        add_paragraph(doc, "`model_loader.py` 管模型加载，`object_detector.py` 管检测与结果对象，`preprocess.py` 管绘图，`web_app.py` 管 HTTP 流程，`index.html` 管展示。模块之间通过明确参数和返回值合作。职责越清楚，修改 OpenCV 样式时越不容易破坏模型加载，替换检测模型时也不必重写整个页面。")

        doc.add_heading("异常与防御式设计", level=2)
        add_paragraph(doc, "CPSC 210 中的健壮性同样体现在前置条件和异常语义。阈值必须位于 0 到 1，模型路径必须存在，图片必须可读取。代码在接近错误来源的位置验证条件，再由 Web 层把不同异常映射为合适的 HTTP 状态。这比在最外层只写一个宽泛异常并返回原始堆栈更容易维护。")

        doc.add_heading("知识关系梳理", level=2)
        add_paragraph(doc, "可以把 CPSC 210 与 AI 项目的结合分成三层。第一层是领域对象：ModelLoader、ObjectDetector 和 DetectionResult 分别表达加载模型、检测目标和保存结果。第二层是应用流程：web_app 接收请求并协调这些对象。第三层是基础设施：PyTorch、Ultralytics、OpenCV 和文件系统提供具体能力。领域层不应该被某个网页按钮绑死，基础设施也不应该决定整个业务流程。")
        add_paragraph(doc, "抽象不是为了让代码看起来复杂，而是把容易变化的部分隔离起来。今天检测实现是 YOLOv8，未来可能换成其他版本或远程模型服务。只要新实现仍能接收图片路径和阈值，并返回 DetectionResult 列表，Web 层的主要流程就可以保持稳定。相反，如果 Web 路由直接读取 Ultralytics 每个张量字段，框架升级会迫使页面流程一起修改，耦合会明显增加。")
        add_paragraph(doc, "实例和值对象也承担不同角色。YOLOv8Detector 实例内部持有已经加载的模型，有身份、状态和可重复调用的行为，适合被缓存并服务多次请求。DetectionResult 只描述某一次检测得到的值，两份坐标、类别和置信度完全相同的结果在业务上可以视为同一个值。值对象可以有格式转换方法，但不应该偷偷加载模型或修改全局状态。")
        add_paragraph(doc, "从 GRASP 的角度看，web_app 承担 Controller 角色，把用户事件转换成应用操作；检测器可以视为负责推理的 Information Expert，因为它最了解模型输出；DetectionResult 具有较高内聚，只保存与一个检测目标直接相关的数据。使用抽象接口也具有 Protected Variations 的思想：用稳定接口保护上层代码，减少第三方库变化带来的影响。")
        add_paragraph(doc, "本项目没有为了套用设计模式而创建大量空类。对于当前课堂规模，模块级 Flask 路由和两个模型缓存已经足够清楚。面向对象设计的目标不是类越多越好，而是让职责、依赖和变化方向容易理解。当 web_app 继续增长时，才有理由进一步拆分 routes、services、schemas 和 config。")

        doc.add_heading("如何顺着代码阅读", level=2)
        add_paragraph(doc, "阅读时先从用户操作进入，而不是从 import 开始逐行背诵。先在 index.html 找到分类按钮和 fetch 请求，再到 web_app.py 找 `/predict` 路由；沿着 `save_upload`、`predict_image`、`get_model` 继续追到 model_loader.py；最后回到 JSON 响应和前端渲染。检测流程则从 `/detect` 追到 YOLOv8Detector、DetectionResult 和 draw_detections。每次只回答三个问题：输入是什么，当前模块负责什么，输出交给谁。")
        add_paragraph(doc, "这种阅读方式能把控制流和依赖方向串起来。遇到不理解的语法时，再回头补 Path、dataclass、继承或类型标注等局部知识，不必在开始前把所有概念一次学完。课堂上画出的架构箭头就是阅读路线，代码只是把每条箭头具体实现出来。")

        doc.add_heading("八  课堂实操流程", level=1)
        add_paragraph(doc, "以下流程用于课堂演示现有作品，不要求先下载完整数据，也不要求重新训练模型。仓库已包含轻量类别目录、分类 checkpoint 和 YOLOv8n 权重。")
        add_code(doc, "conda activate pytorch_env\npython -m pip install -r requirements.txt\npython tests/test_env.py\npython app/web_app.py")
        add_numbered(doc, [
            "浏览器访问 http://127.0.0.1:5000，确认首页和资源状态正常。",
            "选择图像分类，上传宠物图片，观察原图预览和 Top 3。说明结果来自 37 类模型，而不是网页写死。",
            "切换目标检测，保持阈值 0.5 上传包含常见目标的图片，观察目标数量、类别、分数和检测框。",
            "用同一张图片分别设置 0.3 和 0.7，比较框数量。只改变一个变量，记录误检与漏检。",
            "点击重新运行和清空，确认按钮状态、结果区域和图片预览都按预期恢复。",
            "尝试错误格式文件，确认系统给出可读提示而不是无信息的 Internal Server Error。",
        ])
        add_paragraph(doc, "终端中的 Flask development server warning 在本地课堂演示中无需修复。它提醒开发者不要把内置服务器直接部署到公网。正式部署应使用生产 WSGI 服务器，并增加反向代理、日志、认证、限流和监控。")

        doc.add_heading("九  如何判断项目真的可用", level=1)
        add_paragraph(doc, "页面能够打开只证明模板和首页路由工作，不证明真实模型、接口和交互都正确。项目验收应形成证据链：环境检查确认依赖与资源；单元测试检查局部逻辑；Mock 接口测试检查状态码和 JSON；真实图片请求检查权重与推理；浏览器操作检查按钮、滑块和结果渲染；验证集评估检查模型质量。")
        add_code(doc, "python -m unittest discover -s tests -p \"test_*.py\"\npython scripts/evaluate_classifier.py")
        add_paragraph(doc, "当前自动化测试共 9 项通过。完整验证集评估需要本地 `data/oxford_pet_split/val`，当前基准为 90.19%。只用仓库中的每类代表图片可以启动成品，但不能据此重新训练或复现完整评估。")

        doc.add_heading("十  本课总结", level=1)
        add_paragraph(doc, "这套系统的主线可以概括为：数据定义模型能学什么，模型把图片变成结构化结果，Flask 把模型能力封装成接口，OpenCV 把检测坐标画回图片，前端把结果呈现给用户，测试和文档保证其他人能够复现。")
        add_paragraph(doc, "作品最值得展示的不是简单调用两个库，而是把不同职责连接成一个可操作、可诊断的闭环。专业表达必须尊重事实：分类模型做过微调，YOLO 没有在本项目中重新训练；OpenCV 负责处理和绘图，不是主要识别模型；置信度阈值影响推理过滤，不改变权重；Oxford IIIT Pet 是宠物品种数据，不是完整毛色分类数据。")

        doc.add_heading("课堂检查题", level=2)
        add_numbered(doc, [
            "为什么分类和检测不能只用同一个 Top 3 结果表示？",
            "为什么模型缓存后，调整检测阈值仍然能够生效？",
            "DetectionResult 为什么是值对象，而 YOLOv8Detector 为什么是有行为的实例？",
            "为什么前端已经检查文件类型，后端还必须重新检查？",
            "如果要支持金渐层、银渐层和蓝猫，第一步应该修改页面、模型还是数据？为什么？",
        ])

        doc.add_heading("课后复习主线", level=2)
        add_paragraph(doc, "复习时不要先背每个函数。先画出用户、前端、Flask、分类模型、检测模型、OpenCV 和文件系统之间的关系，再进入代码寻找每条箭头由哪个函数实现。只要能够从上传图片一路解释到 JSON 与结果图，并能用 CPSC 210 的概念说明模块为什么这样拆分，就已经掌握了本项目的核心。")

        # Keep all section headers, title, and subtitle explicitly black.
        for paragraph in doc.paragraphs:
            if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(BLACK)

        doc.core_properties.title = "AI 图像识别项目课堂讲义"
        doc.core_properties.subject = "MobileNetV3 YOLOv8 Flask OpenCV 与 CPSC 210 综合讲义"
        doc.core_properties.author = ""
        OUTPUT.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUTPUT)
        print(OUTPUT)


if __name__ == "__main__":
    build()
