from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"C:\Vscode\AI-Image-Recognition - 1")
OUTPUT = PROJECT_ROOT / "courseware" / "lesson6_web_integration_detailed.docx"


BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
TITLE_COLOR = RGBColor(46, 116, 181)
SUBTITLE_COLOR = RGBColor(79, 79, 79)
TEXT_COLOR = RGBColor(0, 0, 0)
MUTED_COLOR = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
NOTE_FILL = "F7FAFC"
TABLE_BORDER = "D9E1EA"


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_run_font(
    run,
    *,
    font: str = BODY_FONT,
    size: int | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_paragraph_border(paragraph, *, top: bool = False, bottom: bool = False, color: str = "D9E1EA") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    def _edge(name: str):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "1")
        edge.set(qn("w:color"), color)
        return edge

    if top:
        p_bdr.append(_edge("top"))
    if bottom:
        p_bdr.append(_edge("bottom"))


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    size: int = 11,
    color: RGBColor = TEXT_COLOR,
    bold: bool = False,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: int = 0,
    after: int = 6,
    line: float = 1.25,
    font: str = BODY_FONT,
    style: str | None = None,
) -> object:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if text:
        run = p.add_run(text)
        set_run_font(run, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_runs_paragraph(
    doc: Document,
    parts: list[tuple[str, dict]],
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: int = 0,
    after: int = 6,
    line: float = 1.25,
) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(run, **opts)
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        set_paragraph_spacing(p, before=18, after=10, line=1.15)
    elif level == 2:
        set_paragraph_spacing(p, before=14, after=7, line=1.15)
    else:
        set_paragraph_spacing(p, before=10, after=5, line=1.15)
    run = p.add_run(text)
    size, color = {1: (16, TITLE_COLOR), 2: (13, TITLE_COLOR), 3: (12, RGBColor(31, 77, 120))}[level]
    set_run_font(run, size=size, color=color, bold=True)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=0, after=0, line=1.0)
        set_paragraph_shading(p, LIGHT_FILL)
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.18)
        fmt.right_indent = Inches(0.18)
        fmt.first_line_indent = Inches(0)
        run = p.add_run(line)
        set_run_font(run, font=CODE_FONT, size=9, color=TEXT_COLOR)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inch: list[float], indent_inch: float = 0.083) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa(w) for w in widths_inch)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa(indent_inch)))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
    for width in widths_inch:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa(width)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inch[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa(widths_inch[idx])))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill: str | None = None) -> None:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(paragraph, before=0, after=3, line=1.15)
                for run in paragraph.runs:
                    set_run_font(run, size=10, color=TEXT_COLOR)
            if r_idx == 0 and header_fill:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=10, color=TEXT_COLOR, bold=True)


def add_table(doc: Document, rows: list[list[str]], widths_inch: list[float], header_fill: str | None = None) -> object:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
    set_table_geometry(table, widths_inch)
    style_table(table, header_fill=header_fill)
    return table


def set_footer(section, text: str) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=9, color=SUBTITLE_COLOR)
    set_paragraph_spacing(p, before=0, after=0, line=1.0)


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [("Heading 1", 16, TITLE_COLOR), ("Heading 2", 13, TITLE_COLOR), ("Heading 3", 12, RGBColor(31, 77, 120))]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color


def build_doc() -> Document:
    doc = Document()
    set_document_styles(doc)
    section = doc.sections[0]
    set_section_geometry(section)
    set_footer(section, "第六课时 · Flask Web界面开发与系统集成 · AI-Image-Recognition")

    # Cover page
    add_paragraph(
        doc,
        "第六课时课件",
        size=24,
        color=TITLE_COLOR,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=4,
        line=1.0,
    )
    add_paragraph(
        doc,
        "第六课时：Flask Web界面开发与系统集成（小白详解版）",
        size=14,
        color=SUBTITLE_COLOR,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=10,
        line=1.15,
    )
    add_paragraph(
        doc,
        "从命令行推理到浏览器上传识别，把 AI 模型接进真正能展示的应用。",
        size=11,
        color=TEXT_COLOR,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=12,
        line=1.25,
    )

    intro_rows = [
        ["项目", "AI-Image-Recognition"],
        ["课程重点", "Flask 路由、文件上传、模型单例、前后端联调"],
        ["相关文件", "app/web_app.py、app/model_loader.py、templates/index.html"],
        ["数据前提", "data/oxford_pet_split/train 和 val 已准备好"],
    ]
    add_table(doc, intro_rows, [1.35, 5.15], header_fill=LIGHT_BLUE_FILL)

    add_paragraph(
        doc,
        "本课的核心不是“写一个花哨网页”，而是把已经训练好的图像识别能力，包装成一个真正能给用户使用的 Web 服务。你可以把它理解成：前面几课是在修发动机，这一课是在给发动机装方向盘、车窗和仪表盘。",
        size=11,
        after=8,
    )
    add_paragraph(
        doc,
        "为了方便小白理解，这份课件会先讲概念，再讲整体流程，最后才逐个拆代码。你不需要一开始就记住所有函数，只要先知道它们各自负责什么，再去看代码，理解会轻很多。",
        size=11,
        after=14,
    )

    doc.add_page_break()

    add_heading(doc, "1. 这一课你要学什么", 1)
    add_paragraph(
        doc,
        "在开始写代码之前，先把目标说清楚。只要你能完整说出“浏览器上传图片后，Flask 接住请求，模型完成推理，页面显示结果”这条链路，就算抓住了本课的主线。",
    )
    goals = [
        ["你要完成的事", "为什么要学"],
        ["理解 Flask 的路由和视图函数", "知道浏览器请求是怎么找到 Python 函数的"],
        ["理解模板渲染和 JSON 响应", "知道页面和接口分别怎么返回数据"],
        ["理解单例模式", "避免每次请求都重新加载模型"],
        ["理解文件上传和图像预处理", "把用户上传的原图变成模型能吃的张量"],
        ["理解前后端联调", "知道页面如何把结果展示出来"],
    ]
    add_table(doc, goals, [1.9, 4.6], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "如果你是第一次接触 Web 开发，不要着急。你只需要先记住一个最朴素的顺序：先上传，再预处理，再推理，再展示。后面的所有代码，都是为了让这四步稳定地跑起来。",
        after=10,
    )

    add_heading(doc, "2. 为什么要加 Web 层", 1)
    add_paragraph(
        doc,
        "前几课我们已经把模型训练出来了，也能在命令行里做识别，但命令行对普通用户并不友好。用户要自己找命令、记路径、切环境，出错概率很高。Web 层的作用，就是把复杂的技术细节藏起来，让用户只做“点按钮、传图片、看结果”这三件事。",
    )
    problems = [
        ["命令行版的问题", "对用户的影响", "Web 层怎么解决"],
        ["需要记住命令和参数", "新手容易输错", "把操作做成按钮和上传框"],
        ["只能在终端里看结果", "不方便展示给别人看", "用浏览器页面展示识别结果"],
        ["每次都手动调用脚本", "交互性很差", "浏览器上传后自动返回结果"],
        ["模型加载过程不透明", "不容易讲解", "在代码里明确展示单例加载流程"],
    ]
    add_table(doc, problems, [1.95, 2.25, 2.3], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "对于课程展示来说，Web 层还有一个额外价值：它能把“模型”变成“产品”。这一步非常重要，因为很多同学第一次做 AI 项目，卡住的不是模型准确率，而是不会把模型接到一个能演示的界面上。",
        after=10,
    )

    add_heading(doc, "3. Flask 到底是什么", 1)
    add_paragraph(
        doc,
        "Flask 是一个轻量级的 Python Web 框架。你可以把它想成一个“接线板”：用户在浏览器里发来请求，Flask 把请求接住，再交给你写好的 Python 函数处理，最后把结果送回浏览器。",
    )
    flask_table = [
        ["概念", "小白解释", "在我们项目里的位置"],
        ["路由 route", "浏览器地址和 Python 函数之间的映射", "`/` 和 `/predict`"],
        ["视图函数 view", "真正处理请求的函数", "`index()` 和 `predict()`"],
        ["模板 template", "页面的 HTML 文件", "`templates/index.html`"],
        ["JSON 响应", "接口返回的结构化数据", "识别结果、图片地址、错误信息"],
    ]
    add_table(doc, flask_table, [1.3, 3.05, 2.15], header_fill=LIGHT_BLUE_FILL)
    add_code_block(
        doc,
        """
from flask import Flask

app = Flask(__name__)

@app.route("/")
def index():
    return "Hello Flask"
""",
    )
    add_paragraph(
        doc,
        "这段最小示例的意思很简单：当用户访问首页时，Flask 会调用 `index()` 函数，并把函数返回的内容显示到浏览器里。等你理解了这一个最小例子，再去看本课的完整版本，就不会觉得陌生。",
        after=10,
    )

    add_heading(doc, "4. 这次项目的目录怎么看", 1)
    add_paragraph(
        doc,
        "在学习 Web 集成时，先把文件职责分清楚。你不用记住所有代码，只需要知道每个文件在哪一层、干什么、为什么要存在。",
    )
    files = [
        ["文件", "职责", "小白记法"],
        ["app/web_app.py", "Flask 入口、上传处理、预测接口", "网站的大门口"],
        ["app/model_loader.py", "加载预训练和微调模型", "把模型从文件里取出来"],
        ["app/transforms.py", "定义训练和验证预处理", "把图片变成模型能吃的样子"],
        ["app/compare_inference.py", "对比预训练和微调模型", "看微调有没有提升"],
        ["templates/index.html", "前端页面", "用户真正看到的界面"],
        ["static/uploads/", "保存用户上传图片的临时目录", "中转站"],
        ["models/oxford_pet_mobilenet_epoch1.pth", "微调后的模型权重", "识别能力的核心文件"],
    ]
    add_table(doc, files, [1.45, 2.7, 2.35], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "在这一课里，最重要的不是多写代码，而是建立“层”的概念。模型、页面、接口、文件保存、预处理，这些东西各自负责一件事，分开以后才容易维护。",
        after=10,
    )

    add_heading(doc, "5. `web_app.py` 一次看懂", 1)
    add_paragraph(
        doc,
        "下面进入最重要的文件：`app/web_app.py`。你可以把它当成网站的总控制器。它做的事情很多，但拆开之后其实只有几类：初始化、加载模型、接收图片、做预测、返回结果。",
    )

    add_heading(doc, "5.1 初始化应用和上传目录", 2)
    add_code_block(
        doc,
        """
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
""",
    )
    add_paragraph(
        doc,
        "第一行创建 Flask 应用，后面两行配置上传目录和最大文件大小，最后一行确保上传目录真的存在。对小白来说，这段代码的意义只有一个：先把“文件要放哪儿”这件事安排好，不然后面上传时程序会找不到地方存图。",
        after=8,
    )

    add_heading(doc, "5.2 怎么找到数据集和类别名", 2)
    add_code_block(
        doc,
        """
@lru_cache(maxsize=1)
def load_class_names() -> tuple[str, ...]:
    train_dir = resolve_data_root() / "train"
    dataset = ImageFolder(train_dir)
    return tuple(dataset.classes)
""",
    )
    add_paragraph(
        doc,
        "这里用的是 `ImageFolder`。它的规则很简单：训练集里每个类别必须是一个文件夹，文件夹名就是类别名。比如 `Abyssinian/`、`British_Shorthair/`、`Russian_Blue/`。程序不需要你额外写标签表，直接读文件夹名就行。",
        after=8,
    )
    add_paragraph(
        doc,
        "我们给 `load_class_names()` 加了 `lru_cache(maxsize=1)`，意思是只加载一次类别名，之后直接复用缓存结果。这样既清晰，又省时间。",
        after=10,
    )

    add_heading(doc, "5.3 为什么模型也要只加载一次", 2)
    add_code_block(
        doc,
        """
_model = None

def get_model():
    global _model
    if _model is None:
        loader = FineTunedLoader(model_path, num_classes=len(class_names))
        _model = loader.load_model()
    return _model
""",
    )
    add_paragraph(
        doc,
        "这就是单例思想的简单版本。第一次有人访问网页时，模型还没有被加载，于是程序就去读权重文件；之后再有人上传图片时，就直接用同一个模型对象，不再重复加载。",
        after=8,
    )
    add_paragraph(
        doc,
        "为什么要这样做？因为模型文件通常不小，加载一次可能要几秒。如果每个请求都重新加载模型，网页会很慢，服务器也会很累。对新手来说，你只要记住一句话：模型不是请求级对象，而是服务级对象。",
        after=10,
    )

    add_heading(doc, "5.4 图像怎么变成张量", 2)
    add_code_block(
        doc,
        """
def preprocess_image(image_path: Path) -> torch.Tensor:
    image = Image.open(image_path).convert("RGB")
    pipeline = transforms.Compose(
        [
            transforms.Resize(256),
            transforms.CenterCrop(224),
            transforms.ToTensor(),
            transforms.Normalize(
                mean=[0.485, 0.456, 0.406],
                std=[0.229, 0.224, 0.225],
            ),
        ]
    )
    return pipeline(image).unsqueeze(0)
""",
    )
    add_paragraph(
        doc,
        "用户上传的是普通图片文件，模型需要的是张量。这个转换过程叫预处理。`Resize(256)` 先把短边缩放到 256，`CenterCrop(224)` 再从中间裁出 224 x 224，`ToTensor()` 把图片转成数值，`Normalize()` 则把数值归一化到和 ImageNet 训练时一致的范围。",
        after=8,
    )
    add_paragraph(
        doc,
        "这里的关键是：训练时怎么预处理，验证和推理时就尽量保持一致。否则模型看到的输入分布不一样，结果就会不稳定。",
        after=10,
    )

    add_heading(doc, "5.5 预测函数怎么写", 2)
    add_code_block(
        doc,
        """
def predict_image(image_path: Path, topk: int = 3) -> list[dict[str, float | str]]:
    model = get_model()
    class_names = load_class_names()
    image_tensor = preprocess_image(image_path)

    with torch.no_grad():
        logits = model(image_tensor)
        probs = F.softmax(logits[0], dim=0)
        top_prob, top_idx = torch.topk(probs, k=min(topk, len(class_names)))

    results: list[dict[str, float | str]] = []
    for prob, idx in zip(top_prob, top_idx):
        results.append(
            {"label": class_names[int(idx)], "score": round(float(prob) * 100.0, 2)}
        )
    return results
""",
    )
    add_paragraph(
        doc,
        "这段代码做了四件事：先取模型，再取类别名，然后预处理图片，最后输出 Top-K 结果。`torch.no_grad()` 的作用是告诉 PyTorch：现在只是推理，不需要算梯度，这样会更省显存、更快。",
        after=8,
    )
    add_paragraph(
        doc,
        "`softmax` 会把模型输出的分数变成概率。`topk` 则会找出概率最高的几个类别。对于课堂演示来说，返回 Top-3 已经足够清楚。",
        after=10,
    )

    add_heading(doc, "5.6 上传文件怎么保存", 2)
    add_code_block(
        doc,
        """
def save_upload(upload) -> Path:
    filename = secure_filename(upload.filename)
    if not filename:
        raise ValueError("Uploaded file name is empty.")

    target = UPLOAD_FOLDER / filename
    upload.save(target)
    return target
""",
    )
    add_paragraph(
        doc,
        "`secure_filename()` 会把上传文件名清理成安全的形式，避免奇怪字符或者路径问题。对小白来说，这一步可以理解成“先把用户给的文件名洗干净，再放进上传目录”。",
        after=10,
    )

    add_heading(doc, "5.7 两个路由分别做什么", 2)
    add_code_block(
        doc,
        """
@app.route("/")
def index():
    return render_template("index.html", ...)

@app.route("/predict", methods=["POST"])
def predict():
    ...
    return jsonify({"results": results, "image_url": image_url})
""",
    )
    add_paragraph(
        doc,
        "首页路由 `/` 负责把页面送给用户；`/predict` 路由负责接收图片、做预测，再把 JSON 结果返回前端。页面和接口分开以后，结构会清楚很多：页面负责好看，接口负责干活。",
        after=10,
    )

    add_heading(doc, "6. `FineTunedLoader` 的作用", 1)
    add_paragraph(
        doc,
        "模型单例能工作，前提是我们能正确把权重文件加载成模型对象。这个工作由 `app/model_loader.py` 里的 `FineTunedLoader` 完成。你可以把它看成“模型装配工”：先搭出模型骨架，再把训练好的权重装进去。",
    )
    add_code_block(
        doc,
        """
class FineTunedLoader(ModelLoader):
    def __init__(self, model_path: str | Path, num_classes: int, freeze_features: bool = True) -> None:
        self.model_path = Path(model_path)
        self.num_classes = num_classes
        self.freeze_features = freeze_features

    def load_model(self) -> nn.Module:
        model = build_finetune_model(
            num_classes=self.num_classes,
            freeze_features=self.freeze_features,
        )
        state_dict = torch.load(self.model_path, map_location="cpu", weights_only=True)
        model.load_state_dict(state_dict)
        model.eval()
        return model
""",
    )
    loader_rows = [
        ["步骤", "做什么", "为什么要这样做"],
        ["构建模型", "先调用 `build_finetune_model()`", "先有结构，后装权重"],
        ["加载权重", "把 checkpoint 读进来", "让模型恢复到训练后的状态"],
        ["切换评估模式", "调用 `model.eval()`", "推理时关闭训练行为"],
    ]
    add_table(doc, loader_rows, [1.1, 2.55, 2.85], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "这里最容易混淆的一点是：`build_finetune_model()` 负责“模型长什么样”，`load_state_dict()` 负责“模型学会了什么”。这两个步骤不是一回事。",
        after=10,
    )

    add_heading(doc, "7. 前端页面到底做了什么", 1)
    add_paragraph(
        doc,
        "前端文件是 `templates/index.html`。它没有使用复杂框架，只有 HTML、CSS 和原生 JavaScript。这样做的好处是，课堂上更容易讲明白，而且同学看代码时不会被额外工具淹没。",
    )
    frontend_rows = [
        ["部分", "职责", "你可以怎么理解"],
        ["HTML", "搭页面骨架", "先把上传框、预览区、结果区放出来"],
        ["CSS", "管样式", "决定页面看起来整不整齐"],
        ["JavaScript", "管交互", "负责上传文件、发请求、更新状态"],
    ]
    add_table(doc, frontend_rows, [1.0, 2.35, 3.15], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "页面上最重要的两个状态是“识别中”和“识别完成”。对用户来说，这两个状态很关键，因为它们告诉用户：系统在工作，不是卡死了。",
        after=8,
    )
    add_code_block(
        doc,
        """
statusBox.textContent = "识别中...";
statusBox.className = "status loading";

fetch("/predict", {
  method: "POST",
  body: formData
})
""",
    )
    add_paragraph(
        doc,
        "`fetch()` 会把图片发给后端，`FormData` 用来模拟表单上传。等后端返回 JSON 之后，JavaScript 再把结果写回页面。整个过程看起来像“网页在和服务器聊天”。",
        after=10,
    )

    add_heading(doc, "8. 前后端联调的完整链路", 1)
    add_paragraph(
        doc,
        "联调这个词听起来复杂，实际上就是“把前端、后端、模型、数据这几块接到一起并逐步排错”。下面这张表是你讲课时最适合顺着说的一条线。",
    )
    flow = [
        ["步骤", "后端做什么", "用户看到什么", "小白理解"],
        ["1", "选择图片并上传", "看到预览图", "图片已经进网页了"],
        ["2", "Flask 接收文件", "显示“识别中...”", "服务器开始干活"],
        ["3", "保存图片到临时目录", "页面不变", "图片先放到一个中转站"],
        ["4", "预处理图片", "页面仍在等待", "把原图变成张量"],
        ["5", "模型推理", "页面仍在等待", "模型开始猜是什么"],
        ["6", "返回 Top-3 JSON", "出现分类结果和百分比", "服务器把答案发回来了"],
        ["7", "前端更新页面", "结果展示完毕", "页面把答案画出来"],
    ]
    add_table(doc, flow, [0.55, 2.05, 1.9, 2.0], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "如果你要用一句话讲清楚联调，就说：前端负责“收图、发请求、展示结果”，后端负责“存图、预处理、推理、返结果”，中间靠 JSON 沟通。",
        after=10,
    )

    add_heading(doc, "9. 你应该怎样运行和验证", 1)
    add_paragraph(
        doc,
        "下面这部分是实操同学最需要的。先把环境装对，再跑服务，再看结果。不要一上来就改代码，否则很容易分不清是环境问题还是逻辑问题。",
    )
    add_code_block(
        doc,
        """
pip install -r requirements.txt
python app/web_app.py
""",
    )
    run_rows = [
        ["命令", "你期待看到什么", "如果没看到怎么办"],
        ["`pip install -r requirements.txt`", "Flask、torch、torchvision 等依赖安装完成", "先检查虚拟环境是不是对的"],
        ["`python app/web_app.py`", "显示 `Running on http://127.0.0.1:5000`", "看是否缺少 Flask 或 checkpoint"],
        ["浏览器打开首页", "出现上传区域", "检查模板文件和路由是否存在"],
        ["上传一张图片", "出现 Top-3 结果", "检查图片预处理和模型加载"],
    ]
    add_table(doc, run_rows, [1.65, 2.6, 2.25], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "如果你在 Windows 上看到一些关于 OpenCL 的报错或临时文件提示，但页面仍然正常工作，这种情况通常可以先不管。真正需要优先处理的是导入失败、权重文件找不到、上传后接口报错这几类问题。",
        after=10,
    )

    add_heading(doc, "10. 常见问题排查", 1)
    issues = [
        ["报错/现象", "最可能的原因", "怎么修"],
        ["`ModuleNotFoundError: torch`", "当前 Python 环境不是装好 PyTorch 的环境", "切到 `pytorch_env` 或重新安装依赖"],
        ["`ModuleNotFoundError: Flask`", "Flask 没装", "执行 `pip install -r requirements.txt`"],
        ["找不到模型权重", "`.pth` 文件不存在或路径不对", "确认 `models/oxford_pet_mobilenet_epoch1.pth` 存在"],
        ["找不到数据集", "`data/oxford_pet_split` 目录不完整", "确认 `train` 和 `val` 都存在"],
        ["上传后没有结果", "前端请求失败或后端报错", "看浏览器控制台和终端日志"],
    ]
    add_table(doc, issues, [2.0, 2.4, 2.1], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "解决问题时记住一个顺序：先看环境，再看路径，再看代码，最后再看页面。很多初学者一开始就盯着前端，其实真正出错的地方往往是环境或文件路径。",
        after=10,
    )

    add_heading(doc, "11. 这节课你真正学会了什么", 1)
    conclusion_rows = [
        ["能力", "你现在已经可以做什么"],
        ["模型能力接入 Web", "把训练好的分类模型放到网页里给人使用"],
        ["单例模式", "让模型只加载一次，减少等待时间"],
        ["文件上传", "接收用户从浏览器传来的图片"],
        ["前后端联调", "把页面、接口、模型、预处理全部接起来"],
        ["工程化思维", "知道一个 AI 项目不只是训练模型，还要能被展示和交互"],
    ]
    add_table(doc, conclusion_rows, [1.75, 4.75], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "如果你已经能解释清楚“浏览器上传图片后发生了什么”，并且能够独立跑起这个 Web 页面，那么这节课的核心目标就达成了。",
        after=8,
    )
    add_paragraph(
        doc,
        "第七课你会继续把这个系统往前推，加入目标检测能力，让它不仅能告诉你“这是什么”，还可以告诉你“它在哪里”。",
        after=10,
    )

    add_heading(doc, "12. 课后作业", 1)
    add_paragraph(
        doc,
        "作业不要只停留在“会跑”上，要尽量把讲解也想清楚。这样你下次面对类似项目时，思路会更稳。",
        after=8,
    )
    homework = [
        ["任务", "具体要求"],
        ["跑通 Web 页面", "启动 Flask，打开页面，上传至少 3 张图片，确认都能返回结果"],
        ["解释单例", "用自己的话说明为什么模型不应该每次请求都重新加载"],
        ["理解状态切换", "说清楚页面为什么需要“识别中 / 成功 / 失败”三种状态"],
        ["扩展思考", "想一想如何让页面支持预训练模型和微调模型切换"],
    ]
    add_table(doc, homework, [1.55, 4.95], header_fill=LIGHT_BLUE_FILL)
    add_paragraph(
        doc,
        "如果你愿意再多做一步，可以给结果区加一个更明显的 Top-1 高亮，或者给上传区加一个“清空图片”按钮。做这些小改动能帮助你更熟悉前后端联调的整个过程。",
        after=10,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(str(OUTPUT))
    print(f"Saved to {OUTPUT}")


if __name__ == "__main__":
    main()
