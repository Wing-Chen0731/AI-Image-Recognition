from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"C:\Vscode\AI-Image-Recognition - 1")
COURSEWARE_DIR = PROJECT_ROOT / "courseware"
OUTPUT_DOCX = COURSEWARE_DIR / "lesson7_detection_integration_detailed.docx"
OUTPUT_MD = COURSEWARE_DIR / "lesson7_detection_integration_detailed.md"

BODY_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"
TITLE_BLUE = RGBColor(31, 111, 235)
HEADING_BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(25, 32, 43)
MUTED = RGBColor(95, 107, 122)
SUCCESS = RGBColor(24, 134, 75)
WARNING = RGBColor(122, 90, 0)
LIGHT_BLUE = "E8F1FF"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E8F7EF"
LIGHT_YELLOW = "FFF7D6"
BORDER = "D9E1EA"


TITLE = "第七课时：把目标检测接入 Web 图像识别系统"
SUBTITLE = "从“判断这张图是什么”到“找出图里有什么、在哪里”"


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_run_font(
    run,
    *,
    font: str = BODY_FONT,
    latin_font: str = LATIN_FONT,
    size: int | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin_font
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_spacing(paragraph, *, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inch: list[float], indent_inch: float = 0.083) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa(width) for width in widths_inch)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa(indent_inch)))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
    for width in widths_inch:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa(width)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_inch[min(idx, len(widths_inch) - 1)]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa(width)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(footer, after=0, line=1.0)
    run = footer.add_run("Lesson 7 - Object Detection Integration")
    set_run_font(run, size=9, color=MUTED)


def add_para(
    doc: Document,
    text: str = "",
    *,
    size: int = 11,
    color: RGBColor = TEXT,
    bold: bool = False,
    italic: bool = False,
    before: int = 0,
    after: int = 6,
    line: float = 1.25,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after, line=line)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: HEADING_BLUE, 2: HEADING_BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    set_spacing(p, after=4, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=4, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def add_callout(doc: Document, title: str, lines: list[str], *, fill: str = LIGHT_BLUE, color: RGBColor = TEXT) -> None:
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2, line=1.2)
    set_paragraph_shading(p, fill)
    run = p.add_run(title)
    set_run_font(run, size=11, color=color, bold=True)
    for line in lines:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=2, line=1.2)
        set_paragraph_shading(p, fill)
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.18)
        fmt.right_indent = Inches(0.18)
        run = p.add_run(line)
        set_run_font(run, size=10, color=color)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=0, line=1.0)
        set_paragraph_shading(p, LIGHT_GRAY)
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.18)
        fmt.right_indent = Inches(0.18)
        run = p.add_run(line)
        set_run_font(run, font=BODY_FONT, latin_font=CODE_FONT, size=9, color=TEXT)
    add_para(doc, "", after=2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], LIGHT_GRAY)
        p = header_cells[idx].paragraphs[0]
        set_spacing(p, after=0, line=1.15)
        run = p.add_run(header)
        set_run_font(run, size=10, color=TEXT, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_spacing(p, after=0, line=1.15)
            run = p.add_run(value)
            set_run_font(run, size=10, color=TEXT)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)


def add_page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_docx() -> None:
    COURSEWARE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=28, after=8, line=1.1)
    run = p.add_run(TITLE)
    set_run_font(run, size=22, color=TITLE_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=18, line=1.1)
    run = p.add_run(SUBTITLE)
    set_run_font(run, size=13, color=MUTED)

    add_callout(
        doc,
        "本节课一句话目标",
        [
            "前面我们已经能让网页判断一张图片属于哪个类别，本节课要让网页进一步画出图片里的目标位置，并返回目标名称、置信度和检测框。",
            "学生最终要完成一个闭环：启动 Flask 页面，上传图片，选择分类或检测，看到模型推理结果，并能解释每个文件在这个闭环中的作用。",
        ],
        fill=LIGHT_GREEN,
        color=SUCCESS,
    )

    add_heading(doc, "一、为什么要上这堂课", 1)
    add_para(doc, "第六课我们已经把图像分类模型接入了网页。分类模型解决的问题是：给它一张图，它告诉我们这张图最像哪个类别。这个能力很重要，但它仍然有明显限制。")
    for item in [
        "如果一张图里有多个物体，分类模型通常只能给出整体类别，不能分别指出每个物体。",
        "如果用户想知道“猫在哪里”“车在哪里”“画面中有几个人”，分类模型回答不了位置问题。",
        "真实系统往往不只需要一个类别名称，还需要位置、数量、置信度和可视化结果。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "所以第七课的主线是从 classification 过渡到 object detection。学生不需要一开始就理解 YOLO 的全部数学细节，先要建立一个清楚的工程认知：分类输出类别，检测输出类别加位置。")

    add_table(
        doc,
        ["能力", "输入", "输出", "适合解决的问题"],
        [
            ["图像分类", "一张图片", "Top-1 或 Top-3 类别和概率", "这张图整体属于什么类别"],
            ["目标检测", "一张图片", "多个检测框、类别、置信度", "图里有什么物体，分别在哪里"],
            ["Web 系统集成", "用户上传图片", "页面结果、JSON、可视化图片", "让普通用户不用写代码也能使用模型"],
        ],
        [1.25, 1.35, 2.0, 1.9],
    )

    add_heading(doc, "二、和前面课程的联系", 1)
    add_table(
        doc,
        ["前面学过的内容", "第七课如何继续使用", "学生要形成的理解"],
        [
            ["数据集 train / val", "分类功能仍然要读取 data/train 或 data/oxford_pet_split/train 的类别名", "数据目录不仅服务训练，也服务网页推理时的类别解释"],
            ["迁移学习和微调", "分类 Tab 使用微调后的 MobileNetV3 权重", "第七课没有推翻第六课，而是在已有分类系统旁边新增检测能力"],
            ["Flask 页面", "继续使用 app/web_app.py 和 templates/index.html", "AI 模型不是孤立脚本，要通过接口接入用户页面"],
            ["CPSC210 抽象思想", "ObjectDetector 是接口，YOLOv8Detector 是实现", "面向接口编程可以让系统以后替换模型更容易"],
            ["错误排查", "继续处理缺文件、环境没装、路径不对、端口占用等问题", "调通项目的能力来自逐层定位问题"],
        ],
        [1.55, 2.25, 2.5],
    )

    add_heading(doc, "三、本节课最终项目结构", 1)
    add_code_block(
        doc,
        r"""
AI-Image-Recognition - 1/
├─ app/
│  ├─ web_app.py              # Flask 后端入口：分类接口 + 检测接口
│  ├─ object_detector.py      # 目标检测抽象和 YOLOv8 实现
│  ├─ preprocess.py           # 检测框绘制和图片后处理
│  └─ model_loader.py         # 第六课已有：分类模型加载
├─ templates/
│  └─ index.html              # 前端页面：分类 Tab + 检测 Tab
├─ static/
│  └─ uploads/                # 保存上传图片和检测后画框图片
├─ data/
│  └─ train/ 或 oxford_pet_split/train/   # 分类类别来源
├─ models/
│  └─ oxford_pet_mobilenet_epoch1.pth     # 分类模型权重，若存在则优先使用
├─ yolov8n.pt                 # YOLOv8 检测权重，首次可自动下载
└─ requirements.txt           # 项目依赖，包含 Flask、torch、ultralytics、opencv-python
""",
    )
    add_callout(
        doc,
        "课堂提醒",
        [
            "分类模型权重和检测模型权重不是同一个文件。MobileNetV3 的 .pth 负责猫狗品种分类，YOLOv8 的 yolov8n.pt 负责通用目标检测。",
            "如果网页首页直接 500，优先看终端 Traceback。常见原因是分类权重文件不存在，而不是 HTML 页面坏了。",
        ],
        fill=LIGHT_YELLOW,
        color=WARNING,
    )

    add_page_break(doc)
    add_heading(doc, "四、核心概念讲解", 1)
    concept_rows = [
        ["分类 Classification", "给整张图贴一个或几个类别标签", "输出 cat、dog、Abyssinian 等类别概率"],
        ["目标检测 Detection", "找出图中的目标位置和类别", "输出 person 0.91，框坐标 x1/y1/x2/y2"],
        ["检测框 Bounding Box", "用矩形框标出物体所在区域", "左上角和右下角两个点就能确定一个框"],
        ["置信度 Confidence", "模型对结果有多确定", "0.82 表示模型认为这个检测结果可信度较高"],
        ["阈值 Threshold", "过滤低置信度结果的门槛", "阈值 0.5 表示只保留置信度不低于 0.5 的框"],
        ["Top-K", "分类里展示概率最高的前 K 个类别", "Top-3 表示展示前三个最可能类别"],
        ["JSON", "前后端传输结构化结果的格式", "后端返回 detections，前端用 JS 渲染结果"],
    ]
    add_table(doc, ["概念", "小白解释", "在项目中的样子"], concept_rows, [1.6, 2.35, 2.55])

    add_heading(doc, "五、本节课的总流程图", 1)
    add_code_block(
        doc,
        r"""
用户打开浏览器
    ↓
Flask 返回 templates/index.html
    ↓
用户选择“图像分类”或“目标检测”
    ↓
前端 JavaScript 把图片上传到 /predict 或 /detect
    ↓
后端保存图片到 static/uploads
    ↓
分类路径：MobileNetV3 输出 Top-3 类别
检测路径：YOLOv8 输出检测框、类别、置信度
    ↓
后端返回 JSON
    ↓
前端展示结果；检测路径额外展示画框后的图片
""",
    )

    add_heading(doc, "六、每个代码文件是什么、作用是什么", 1)
    add_heading(doc, "1. app/object_detector.py", 2)
    add_para(doc, "这个文件负责把“目标检测”这件事封装成一个清楚的模块。它不是网页，也不是画图工具，而是检测能力的核心入口。")
    add_table(
        doc,
        ["代码块", "是什么", "作用和目的"],
        [
            ["DetectionResult", "一个 dataclass 值对象", "保存一个检测结果：框坐标、类别名、置信度。to_dict() 把 Python 对象转成前端容易读取的字典。"],
            ["ObjectDetector", "抽象接口", "规定所有检测器都必须有 detect(image_path) 方法。以后换成别的检测模型时，Web 层不用大改。"],
            ["YOLOv8Detector", "YOLOv8 的具体实现", "加载 ultralytics 的 YOLO 模型，读取图片，执行检测，把 YOLO 原始结果转成 DetectionResult。"],
            ["YOLO_CONFIG_DIR", "Ultralytics 配置目录", "把配置写到项目里的 .ultralytics，减少不同电脑上权限或用户目录差异带来的问题。"],
        ],
        [1.35, 1.45, 3.7],
    )
    add_code_block(
        doc,
        r"""
detector = YOLOv8Detector(model_path="yolov8n.pt", conf_threshold=0.5)
detections = detector.detect("static/uploads/demo.jpg")
""",
    )
    add_para(doc, "课堂讲法：这里可以把 YOLOv8Detector 理解成一个“检测工人”。你给它图片路径，它返回若干张小纸条，每张纸条记录一个物体的位置、名字和可信度。")

    add_heading(doc, "2. app/preprocess.py", 2)
    add_para(doc, "这个文件在第七课里承担“后处理”的角色。YOLO 检测完以后只有数字坐标，用户看数字很难理解，所以我们需要把框画回原图。")
    for item in [
        "读取原图：使用 OpenCV 的 cv2.imread()。",
        "遍历检测结果：每一个 DetectionResult 对应一个矩形框。",
        "画矩形框：使用 cv2.rectangle()。",
        "写类别和置信度：使用 cv2.putText()。",
        "保存新图片：生成 xxx_detected.jpg，前端展示这张带框图片。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "为什么要单独放在 preprocess.py",
        [
            "因为检测模型只负责“算结果”，图片绘制负责“展示结果”。把两件事分开，代码更清楚，也更符合单一职责原则。",
        ],
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "3. app/web_app.py", 2)
    add_para(doc, "这是网页系统的后端入口，也是第六课和第七课连接最紧密的文件。第六课的 /predict 仍然保留，第七课新增 /detect。")
    add_table(
        doc,
        ["函数或变量", "作用", "课堂解释"],
        [
            ["resolve_data_root()", "找到分类数据集目录", "为了读取类别名，页面要知道分类模型一共有多少类。"],
            ["resolve_model_path()", "找到微调后的分类权重", "没有分类权重时首页可能直接报 500，这是常见复现问题。"],
            ["get_model()", "只加载一次分类模型", "模型加载很慢，不能每次上传都重新加载。"],
            ["get_detector()", "只加载一次检测模型", "阈值不变时复用 YOLOv8Detector，提高响应速度。"],
            ["predict_image()", "执行分类推理", "把图片变成 tensor，输入 MobileNetV3，输出 Top-3。"],
            ["detect_image()", "执行检测推理", "调用 YOLOv8Detector，并调用 draw_detections 生成带框图片。"],
            ["/predict", "分类接口", "前端选择分类 Tab 时调用它。"],
            ["/detect", "检测接口", "前端选择检测 Tab 时调用它。"],
        ],
        [1.55, 2.05, 2.9],
    )
    add_code_block(
        doc,
        r"""
@app.route("/detect", methods=["POST"])
def detect():
    conf_threshold = parse_conf_threshold()
    image_path = save_upload(file)
    detections, rendered_path = detect_image(image_path, conf_threshold)
    return jsonify({
        "detections": detections,
        "image_url": image_url(rendered_path),
    })
""",
    )
    add_para(doc, "这段代码可以用一句话讲给学生：浏览器把图片传给 /detect，后端保存图片，YOLO 找目标，OpenCV 画框，最后把结果打包成 JSON 还给浏览器。")

    add_heading(doc, "4. templates/index.html", 2)
    add_para(doc, "这是用户看到的页面。它不直接做 AI 推理，只负责交互：选择图片、预览图片、选择模式、发送请求、展示结果。")
    add_table(
        doc,
        ["前端区域", "作用", "学生需要理解"],
        [
            ["tabs", "分类和检测两个按钮", "同一个页面可以调用不同后端接口。"],
            ["dropzone / file-input", "选择或拖入图片", "用户操作最终会变成一个文件对象。"],
            ["threshold slider", "检测置信度阈值滑块", "阈值低结果多，阈值高结果少。"],
            ["fetch('/predict')", "调用分类接口", "前端通过 HTTP 请求和 Flask 通信。"],
            ["fetch('/detect')", "调用检测接口", "检测需要额外提交 conf_threshold。"],
            ["renderClassification()", "渲染 Top-3 分类结果", "把 JSON 里的 label 和 score 显示成结果卡片。"],
            ["renderDetections()", "渲染检测结果", "展示带框图片、类别、置信度和框坐标。"],
        ],
        [1.55, 2.05, 2.9],
    )

    add_heading(doc, "5. requirements.txt", 2)
    add_para(doc, "这个文件记录项目需要安装哪些第三方库。第七课新增和检测强相关的库是 ultralytics，画框依赖 opencv-python。")
    add_code_block(
        doc,
        r"""
torch
torchvision
opencv-python
Pillow
Flask
ultralytics
""",
    )
    add_callout(
        doc,
        "学生常见误区",
        [
            "`pip install -r requirements.txt` 不是启动项目，它只是安装项目依赖。",
            "必须先激活正确 conda 环境，再安装依赖和运行项目。否则包可能装到 base 环境，课堂环境仍然找不到。",
        ],
        fill=LIGHT_YELLOW,
        color=WARNING,
    )

    add_page_break(doc)
    add_heading(doc, "七、课堂详细操作步骤", 1)
    steps = [
        ("确认打开的是项目根目录", "在 VS Code 中打开 AI-Image-Recognition - 1，终端路径应该在项目根目录。不要只打开 app 文件夹。"),
        ("激活 conda 环境", "运行 conda activate pytorch_env。如果环境名不同，就用自己创建的环境名。"),
        ("安装依赖", "运行 pip install -r requirements.txt。第七课需要 ultralytics，如果没有安装，检测功能会报 ModuleNotFoundError。"),
        ("确认分类模型权重存在", "检查 models/oxford_pet_mobilenet_epoch1.pth 或 finetuned_mobilenet.pth 是否存在。没有它时，分类首页会因为找不到 checkpoint 报 500。"),
        ("确认数据目录存在", "检查 data/train 或 data/oxford_pet_split/train 是否存在。它用来读取分类类别名。"),
        ("启动 Web 项目", "运行 python app/web_app.py，看到 Running on http://127.0.0.1:5000 后打开浏览器。"),
        ("测试分类功能", "在图像分类 Tab 上传猫狗图片，观察 Top-3 结果。"),
        ("测试检测功能", "切换到目标检测 Tab，上传包含人、猫、狗、车等常见物体的图片。第一次运行可能会下载 yolov8n.pt。"),
        ("调节阈值", "把阈值从 0.5 改到 0.3 和 0.7，观察检测框数量变化。"),
        ("总结闭环", "让学生用自己的话复述：前端上传，后端保存，模型推理，后端返回 JSON，前端展示。"),
    ]
    for idx, (title, detail) in enumerate(steps, start=1):
        add_heading(doc, f"步骤 {idx}：{title}", 3)
        add_para(doc, detail)
        if idx == 2:
            add_code_block(doc, "conda activate pytorch_env")
        elif idx == 3:
            add_code_block(doc, "pip install -r requirements.txt")
        elif idx == 6:
            add_code_block(doc, "python app/web_app.py")

    add_heading(doc, "八、课堂演示建议", 1)
    add_table(
        doc,
        ["时间", "教师操作", "学生观察点"],
        [
            ["0-10 分钟", "回顾第六课分类系统", "分类只能回答整体类别，不能回答位置。"],
            ["10-25 分钟", "讲目标检测基本概念", "框、类别、置信度、阈值分别是什么。"],
            ["25-45 分钟", "讲 object_detector.py 和 preprocess.py", "检测和画框是两件事。"],
            ["45-65 分钟", "讲 web_app.py 的 /predict 与 /detect", "同一个 Web 后端可以接入多种 AI 能力。"],
            ["65-85 分钟", "学生跟跑完整项目", "能启动、上传、查看结果。"],
            ["85-100 分钟", "故意调阈值和制造小错误", "训练学生用 Traceback 定位问题。"],
            ["100-120 分钟", "总结和布置作业", "用自己的话解释完整闭环。"],
        ],
        [1.1, 2.75, 2.65],
    )

    add_heading(doc, "九、常见问题和解决方案", 1)
    add_table(
        doc,
        ["现象", "常见原因", "解决方式"],
        [
            ["Internal Server Error", "后端抛异常，浏览器只看到 500", "回到终端看 Traceback 最后一段。不要只看浏览器页面。"],
            ["No fine-tuned checkpoint found", "分类模型权重不存在", "把训练好的 .pth 放到 models/oxford_pet_mobilenet_epoch1.pth 或项目根目录 finetuned_mobilenet.pth。"],
            ["No dataset split found", "data/train 不存在，或数据放错层级", "保证存在 data/train/cat、data/train/dog，或 data/oxford_pet_split/train。"],
            ["ModuleNotFoundError: ultralytics", "没有在当前 conda 环境安装依赖", "激活环境后重新运行 pip install -r requirements.txt。"],
            ["第一次检测很慢", "正在下载 yolov8n.pt 或首次加载模型", "等待下载完成；网络不稳定时可提前把 yolov8n.pt 放到项目根目录。"],
            ["检测不到目标", "阈值太高、图片物体不在 COCO 类别中、图片太模糊", "先把阈值调到 0.3，再换常见目标图片测试。"],
            ["TemplateNotFound: index.html", "Flask 没找到 templates 目录", "确认从项目根目录运行 python app/web_app.py，且 templates/index.html 存在。"],
            ["中文显示乱码", "文件编码或终端编码不一致", "优先不影响功能；后续可统一保存为 UTF-8 并重启编辑器。"],
            ["端口 5000 被占用", "已有 Flask 服务没关", "关闭旧终端，或运行 python app/web_app.py --port 5001。"],
        ],
        [1.75, 2.1, 2.65],
    )

    add_heading(doc, "十、给小白学生的理解路线", 1)
    add_code_block(
        doc,
        r"""
先理解用户流程：
上传图片 → 选择分类/检测 → 页面展示结果

再理解系统流程：
HTML 页面 → Flask 接口 → 模型推理 → JSON → 页面渲染

最后理解代码结构：
web_app.py 管流程
object_detector.py 管检测模型
preprocess.py 管画框图片
index.html 管用户界面
requirements.txt 管依赖安装
""",
    )
    for item in [
        "不要一开始就钻 YOLO 的论文细节。课堂目标是先把检测能力接进项目。",
        "先会跑，再会解释，再会修改。顺序不能反。",
        "看到报错先读最后三行，找关键词：FileNotFoundError、ModuleNotFoundError、TemplateNotFound。",
        "把分类和检测放在一张表里比较，学生会更容易建立概念边界。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十一、课堂提问设计", 1)
    for question in [
        "为什么分类模型不能告诉我们猫在图片的哪个位置？",
        "检测结果里的 x1、y1、x2、y2 分别代表什么？",
        "为什么置信度阈值调低以后，检测框通常会变多？",
        "为什么模型加载要做成单例，而不是每次请求都重新加载？",
        "如果 /detect 报 ultralytics not installed，应该检查哪个文件和哪一步操作？",
        "为什么第七课没有删除第六课的 /predict，而是在旁边新增 /detect？",
    ]:
        add_bullet(doc, question)

    add_heading(doc, "十二、课后作业", 1)
    add_table(
        doc,
        ["作业", "要求", "验收标准"],
        [
            ["基础复现", "在自己的电脑上启动项目，分别测试分类和检测", "提交一张分类结果截图和一张检测结果截图。"],
            ["阈值观察", "用同一张图片测试 0.3、0.5、0.7 三个阈值", "写出检测框数量如何变化，并解释原因。"],
            ["代码解释", "用自己的话解释 /detect 的执行过程", "必须包含上传、保存、YOLO 推理、画框、返回 JSON。"],
            ["错误排查", "故意把 yolov8n.pt 移走或卸载 ultralytics，观察报错", "能说出报错原因和恢复方式。"],
            ["拓展挑战", "尝试把检测结果按置信度从高到低展示", "页面结果顺序正确，且不影响原有分类功能。"],
        ],
        [1.35, 3.0, 2.15],
    )

    add_heading(doc, "十三、本节课收尾总结", 1)
    add_para(doc, "第七课的重点不是单纯认识一个新模型，而是让学生看到 AI 项目从模型脚本走向完整应用的过程。")
    for item in [
        "第六课：分类模型接入 Web，解决“这张图是什么”。",
        "第七课：检测模型接入 Web，解决“图里有什么、在哪里”。",
        "工程闭环：前端上传图片，后端接收请求，模型完成推理，结果通过 JSON 回到前端，页面展示给用户。",
        "代码组织：分类、检测、画框、页面渲染各司其职，后续才有继续扩展的空间。",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "下一节课可以自然过渡到",
        [
            "模型效果评估：检测结果怎么判断好不好。",
            "项目优化：错误提示中文化、日志记录、接口稳定性。",
            "部署思路：开发服务器和生产服务器有什么区别。",
        ],
        fill=LIGHT_GREEN,
        color=SUCCESS,
    )

    doc.save(OUTPUT_DOCX)


def build_markdown() -> None:
    content = f"""# {TITLE}

{SUBTITLE}

## 本节课一句话目标

前面我们已经能让网页判断一张图片属于哪个类别。本节课要让网页进一步画出图片里的目标位置，并返回目标名称、置信度和检测框。

学生最终要完成一个闭环：启动 Flask 页面，上传图片，选择分类或检测，看到模型推理结果，并能解释每个文件在这个闭环中的作用。

## 一、为什么要上这堂课

第六课我们已经把图像分类模型接入了网页。分类模型解决的问题是：给它一张图，它告诉我们这张图最像哪个类别。但真实应用里经常还要回答：

- 图里有几个物体？
- 物体分别在哪里？
- 哪些目标可信度高，哪些只是模型猜测？
- 能不能把结果直接画在图片上给用户看？

所以第七课从 `classification` 过渡到 `object detection`。一句话区别：

| 能力 | 解决的问题 | 输出 |
|---|---|---|
| 图像分类 | 这张图整体是什么 | 类别和概率 |
| 目标检测 | 图里有什么，分别在哪里 | 检测框、类别、置信度 |

## 二、和前面课程的联系

| 前面内容 | 第七课怎么用 | 学生要理解 |
|---|---|---|
| train / val 数据集 | 分类功能仍然读取类别名 | 数据目录不仅用于训练，也用于推理解释 |
| 微调 MobileNetV3 | 分类 Tab 继续使用第六课模型 | 第七课是在已有系统旁边新增检测能力 |
| Flask Web | 继续使用 `app/web_app.py` | AI 模型要通过接口给用户使用 |
| CPSC210 抽象 | `ObjectDetector` 是接口，`YOLOv8Detector` 是实现 | 面向接口编程方便以后换模型 |

## 三、项目结构

```text
AI-Image-Recognition - 1/
├─ app/
│  ├─ web_app.py
│  ├─ object_detector.py
│  ├─ preprocess.py
│  └─ model_loader.py
├─ templates/
│  └─ index.html
├─ static/uploads/
├─ data/train/ 或 data/oxford_pet_split/train/
├─ models/oxford_pet_mobilenet_epoch1.pth
├─ yolov8n.pt
└─ requirements.txt
```

## 四、核心概念

- 分类：给整张图贴标签。
- 检测：找出图中每个目标的位置和类别。
- 检测框：用矩形框标出物体区域。
- 置信度：模型对检测结果的确定程度。
- 阈值：过滤低置信度检测结果的门槛。
- JSON：后端返回给前端的结构化结果。

## 五、课堂操作步骤

1. 打开项目根目录。
2. 激活 conda 环境。

```bash
conda activate pytorch_env
```

3. 安装依赖。

```bash
pip install -r requirements.txt
```

4. 确认分类权重存在：`models/oxford_pet_mobilenet_epoch1.pth` 或 `finetuned_mobilenet.pth`。
5. 确认数据目录存在：`data/train` 或 `data/oxford_pet_split/train`。
6. 启动 Flask。

```bash
python app/web_app.py
```

7. 浏览器打开 `http://127.0.0.1:5000`。
8. 测试分类 Tab，观察 Top-3。
9. 测试检测 Tab，观察检测框、类别和置信度。
10. 调整阈值 0.3、0.5、0.7，观察检测框数量变化。

## 六、每个文件的作用

### `app/object_detector.py`

- `DetectionResult`：保存一个检测结果，包括框坐标、类别、置信度。
- `ObjectDetector`：检测器接口，规定必须有 `detect()` 方法。
- `YOLOv8Detector`：YOLOv8 的具体实现，负责加载 `yolov8n.pt` 并执行检测。

### `app/preprocess.py`

- 读取原图。
- 根据检测结果画矩形框。
- 写上类别和置信度。
- 保存带框图片给网页展示。

### `app/web_app.py`

- `/predict`：第六课分类接口。
- `/detect`：第七课检测接口。
- `get_model()`：分类模型只加载一次。
- `get_detector()`：检测模型只加载一次。
- `detect_image()`：检测并生成带框图片。

### `templates/index.html`

- 分类和检测两个 Tab。
- 图片上传和预览。
- 阈值滑块。
- 使用 `fetch()` 调用 Flask 接口。
- 把 JSON 结果渲染到页面。

## 七、常见问题

| 现象 | 原因 | 解决 |
|---|---|---|
| Internal Server Error | 后端异常 | 看终端 Traceback 最后三行 |
| No fine-tuned checkpoint found | 分类权重缺失 | 放入 `.pth` 权重文件 |
| ModuleNotFoundError: ultralytics | 依赖没装到当前环境 | 激活环境后重新 `pip install -r requirements.txt` |
| 第一次检测慢 | 首次下载或加载 YOLOv8 | 等待下载完成，或提前放好 `yolov8n.pt` |
| TemplateNotFound | 模板路径不对或文件缺失 | 确认从项目根目录运行，且 `templates/index.html` 存在 |

## 八、课后作业

| 作业 | 要求 |
|---|---|
| 基础复现 | 分别提交分类截图和检测截图 |
| 阈值观察 | 用同一张图测试 0.3、0.5、0.7 |
| 代码解释 | 用自己的话解释 `/detect` 完整流程 |
| 错误排查 | 解释一个 500 错误的原因和解决方案 |
| 拓展挑战 | 尝试按置信度排序检测结果 |

## 九、本节课总结

第六课解决“这张图是什么”，第七课解决“图里有什么、在哪里”。学生要掌握的不只是 YOLOv8，而是一个完整 AI Web 系统如何把模型能力交给用户使用。
"""
    COURSEWARE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(content, encoding="utf-8")


def main() -> None:
    build_markdown()
    build_docx()
    print(f"Generated: {OUTPUT_MD}")
    print(f"Generated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
