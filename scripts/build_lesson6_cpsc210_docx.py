"""Build a CPSC210 mapping document for lesson 6.

This document extracts the software design ideas from the Flask web integration
lesson and maps them to the current project files.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "courseware" / "lesson6_cpsc210_mapping.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "B8C5D6"


def set_font(run, name: str = "Calibri", size: int | None = None, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        tag = "w:" + margin_name
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in_inches: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width * 1440) for width in widths_in_inches)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = int(widths_in_inches[idx] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = Inches(widths_in_inches[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def style_table(table, widths: list[float], header_fill: str = LIGHT_BLUE):
    set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_shading(cell, header_fill if row_index == 0 else WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=10, bold=row_index == 0)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("第六课时 CPSC210 关联知识单独讲义")
    set_font(run, size=22, bold=True)
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("主题：从 Flask 图像识别 Web 项目理解面向对象设计、职责划分和系统集成")
    set_font(run, size=11)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True)
        run.font.color.rgb = BLUE if level == 1 else DARK_BLUE
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_font(run, size=11)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_font(run, size=11)


def add_code(doc: Document, lines: list[str]):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(32, 32, 32)


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5], header_fill=LIGHT_GRAY)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_font(run, size=11, bold=True)
    run.font.color.rgb = DARK_BLUE
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_font(run, size=10)
    doc.add_paragraph()


def add_mapping_table(doc: Document):
    rows = [
        ("CPSC210 概念", "本项目对应位置", "课堂讲法"),
        ("职责划分 / 高内聚", "app/web_app.py、app/model_loader.py、templates/index.html", "一个文件只承担一类主要责任：页面展示、请求控制、模型加载分开讲。"),
        ("低耦合", "前端 fetch('/predict')，后端返回 JSON", "前端不需要知道 PyTorch 怎么推理，只需要读 JSON。"),
        ("抽象与接口", "ModelLoader 抽象类，MobileNetV3Loader / FineTunedLoader", "不同模型加载方式遵守同一个 load_model() 约定。"),
        ("前置条件", "resolve_data_root()、resolve_model_path()、save_upload()", "程序先检查数据集、权重文件、上传文件是否存在，再继续运行。"),
        ("异常处理", "try / except、FileNotFoundError、ModelDownloadError", "错误要在合适的位置被发现，并返回能理解的信息。"),
        ("状态管理", "_model、_metadata、lru_cache", "模型和类别名属于服务状态，不应该每次请求都重新加载。"),
        ("MVC / 分层设计", "Model、View、Controller 三层", "Model 负责能力，View 负责展示，Controller 负责组织流程。"),
        ("测试与可验证性", "py_compile、Flask test_client、上传图片测试", "不只看能不能启动，还要验证首页、接口、模型推理。"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
    style_table(table, [1.55, 2.2, 2.75])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "这份文档怎么用",
        "它不是新的 Flask 教程，而是把第六课时中和 CPSC210 相关的软件工程思想单独拎出来。讲课时可以先跑通 Web 页面，再用这份文档解释为什么代码要这样拆、这样命名、这样处理错误。",
    )

    add_heading(doc, "1. 本节课和 CPSC210 的关系", 1)
    add_body(
        doc,
        "第六课的表面目标是做一个 Flask 网页：上传图片，调用模型，返回分类结果。更深一层看，它其实是在训练学生把一个能跑的脚本整理成一个有边界、有职责、有接口的应用系统。这正是 CPSC210 很强调的面向对象设计和工程化思维。",
    )
    add_bullets(
        doc,
        [
            "不是只看 AI 模型准确率，而是看模型能力如何被封装成服务。",
            "不是只写一个长脚本，而是把页面、控制器、模型加载、推理逻辑拆开。",
            "不是只让程序在正常路径能跑，而是提前处理数据集、权重、上传文件等异常情况。",
            "不是只会改代码，而是能说明每个模块的职责和依赖关系。",
        ],
    )

    add_heading(doc, "2. CPSC210 知识点总览", 1)
    add_mapping_table(doc)

    add_heading(doc, "3. 职责划分：为什么要拆文件", 1)
    add_body(
        doc,
        "CPSC210 里经常会强调 single responsibility，也就是一个类或模块最好只负责一类事情。第六课项目里，最重要的拆分是 Model、View、Controller 三层。这个拆分不是为了显得复杂，而是为了让后续修改更容易。",
    )
    add_bullets(
        doc,
        [
            "templates/index.html 负责页面结构、上传交互、结果展示。",
            "app/web_app.py 负责接收浏览器请求、组织数据、返回 HTML 或 JSON。",
            "app/model_loader.py 负责构建模型、加载权重、切换到推理模式。",
            "app/transforms.py 和 web_app.py 里的 preprocess_image() 负责把图片变成模型能处理的张量。",
        ],
    )
    add_callout(
        doc,
        "给小白的讲法",
        "如果所有代码都写在一个文件里，第一次可能能跑，但后面要改页面、换模型、查错误都会很痛苦。拆文件的目的不是为了多写文件，而是为了让每个文件的责任更清楚。",
    )

    add_heading(doc, "4. 抽象与接口：ModelLoader 为什么有价值", 1)
    add_body(
        doc,
        "在 app/model_loader.py 中，ModelLoader 是一个抽象类，它规定子类必须实现 load_model()。这对应 CPSC210 的 abstraction 和 interface 思想：调用者只关心“我能拿到一个可用模型”，不关心模型到底是 MobileNetV3、ResNet18，还是已经微调过的 checkpoint。",
    )
    add_code(
        doc,
        [
            "class ModelLoader(ABC):",
            "    @abstractmethod",
            "    def load_model(self) -> nn.Module:",
            "        \"\"\"Load and return a torch.nn.Module.\"\"\"",
        ],
    )
    add_body(
        doc,
        "有了这个抽象之后，FineTunedLoader 和 MobileNetV3Loader 都可以被看成 ModelLoader 的一种实现。将来如果要支持新的模型，只要新增一个 loader，并保持 load_model() 这个接口不变，其他调用代码就可以少改。",
    )

    add_heading(doc, "5. 单例与缓存：为什么模型只加载一次", 1)
    add_body(
        doc,
        "Web 服务和命令行脚本不同。命令行脚本运行一次就结束，Web 服务会持续接收很多请求。模型权重通常比较大，如果每次上传图片都重新加载模型，速度会很慢，内存也会浪费。因此第六课使用 _model 保存已经加载好的模型对象。",
    )
    add_code(
        doc,
        [
            "_model = None",
            "",
            "def get_model():",
            "    global _model",
            "    if _model is None:",
            "        loader = FineTunedLoader(model_path, num_classes=len(class_names))",
            "        _model = loader.load_model()",
            "    return _model",
        ],
    )
    add_body(
        doc,
        "这不是严格意义上完整的设计模式实现，但它体现了单例思想：在一个服务进程里，模型对象只创建一次，后续请求复用同一个对象。load_class_names() 上的 lru_cache 也是同类思路：类别名只需要从训练目录读取一次。",
    )

    add_heading(doc, "6. 前置条件与异常处理：先检查，再运行", 1)
    add_body(
        doc,
        "CPSC210 会强调 precondition、exception 和 defensive programming。第六课代码里有几个典型检查点：数据集目录必须存在，模型权重必须存在，上传文件名不能为空，类别数至少要大于等于 2。这些检查让错误更早暴露，也让报错更容易定位。",
    )
    add_bullets(
        doc,
        [
            "resolve_data_root() 检查 data/oxford_pet_split/train 或 data/train 是否存在。",
            "resolve_model_path() 检查微调权重文件是否存在。",
            "save_upload() 检查上传文件名是否为空，并保存到 static/uploads。",
            "build_finetune_model() 检查 num_classes 是否至少为 2。",
        ],
    )
    add_callout(
        doc,
        "课堂提醒",
        "不要把异常处理讲成“把报错藏起来”。好的异常处理是让错误更清楚、更靠近真正原因，而不是让程序假装没问题。",
    )

    add_heading(doc, "7. 低耦合：前端为什么只拿 JSON", 1)
    add_body(
        doc,
        "第六课的前后端通过 /predict 这个接口通信。前端用 fetch() 上传 FormData，后端返回 JSON。这样前端不需要理解 PyTorch，也不需要知道模型文件在哪里。后端也不需要拼复杂 HTML，只要返回结构化结果。",
    )
    add_code(
        doc,
        [
            "fetch('/predict', {",
            "  method: 'POST',",
            "  body: formData",
            "})",
            "",
            "return jsonify({",
            "    'results': results,",
            "    'image_url': image_url",
            "})",
        ],
    )
    add_body(
        doc,
        "这就是低耦合的好处：如果以后把页面换成 Vue、React 或小程序，只要接口格式不变，后端模型推理部分可以继续复用。",
    )

    add_heading(doc, "8. 状态设计：页面也有状态", 1)
    add_body(
        doc,
        "CPSC210 不只关心类，也关心对象状态是否清楚。第六课的前端页面虽然没有复杂框架，但已经有明显的状态变化：空闲、选择图片、识别中、识别成功、识别失败。学生要能说明页面为什么要显示不同状态。",
    )
    add_numbered(
        doc,
        [
            "空闲状态：还没有选择图片，等待用户上传。",
            "预览状态：用户已经选择图片，可以看到图片内容。",
            "识别中状态：请求已经发给后端，按钮或提示应显示加载中。",
            "成功状态：后端返回 Top-3 分类结果，页面展示类别和置信度。",
            "失败状态：后端返回错误或网络失败，页面显示错误提示。",
        ],
    )

    add_heading(doc, "9. 可测试性：本节课应该怎么验收", 1)
    add_body(
        doc,
        "从 CPSC210 角度看，验收不能只停留在“我打开网页看起来没问题”。更好的检查方式是按边界逐个验证：语法、首页、接口、模型、错误路径。",
    )
    table = doc.add_table(rows=5, cols=3)
    data = [
        ("检查项", "命令或操作", "通过标准"),
        ("语法检查", "python -m py_compile app\\web_app.py", "没有语法错误。"),
        ("首页检查", "访问 http://127.0.0.1:5000", "返回页面，不再出现 TemplateNotFound。"),
        ("接口检查", "上传一张图片", "返回 Top-3 JSON，页面能渲染结果。"),
        ("错误路径", "不上传文件或上传无效文件", "接口返回清楚的错误信息。"),
    ]
    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    style_table(table, [1.35, 2.7, 2.45])

    add_heading(doc, "10. 可以布置给学生的 CPSC210 小作业", 1)
    add_body(
        doc,
        "下面这些作业不要求学生重新训练模型，重点是让他们把软件设计讲清楚。",
    )
    add_numbered(
        doc,
        [
            "画出第六课项目的 MVC 分层图，并标出每个文件属于哪一层。",
            "解释 ModelLoader 抽象类的作用，并说明 FineTunedLoader 为什么是它的一种实现。",
            "在 get_model() 里加入一行日志，刷新页面多次，证明模型只加载一次。",
            "找出三个前置条件检查，并说明如果没有这些检查会出现什么问题。",
            "给 /predict 接口设计一个错误测试：没有上传文件时应该返回什么。",
            "写一段话说明：为什么前端拿 JSON 比后端直接拼完整 HTML 更容易扩展。",
        ],
    )

    add_heading(doc, "11. 教师讲课顺序建议", 1)
    add_numbered(
        doc,
        [
            "先让学生运行 Web 页面，确认他们知道这个系统最终能做什么。",
            "再回到项目目录，说明每个新增文件的职责。",
            "讲 app/web_app.py 时重点讲 Controller：接请求、调模型、回结果。",
            "讲 app/model_loader.py 时重点讲 abstraction：不同 loader 都实现 load_model()。",
            "讲 get_model() 时重点讲服务状态：模型对象不是每个请求都新建。",
            "最后用测试清单收尾，让学生知道工程代码必须可验证。",
        ],
    )

    add_heading(doc, "12. 一句话总结", 1)
    add_body(
        doc,
        "第六课和 CPSC210 的核心连接点是：把能跑的 AI 脚本整理成一个职责清楚、接口清楚、状态清楚、错误处理清楚、可以测试的应用系统。Flask 只是载体，真正重要的是软件设计思维。",
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
